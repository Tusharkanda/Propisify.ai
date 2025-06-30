from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os

def export_to_word(sections: dict, client_name: str) -> str:
    """
    Creates a professional Word document for the proposal and returns the file path.
    Args:
        sections (dict): Dictionary with keys 'executive_summary', 'scope_of_work', 'pricing'
        client_name (str): Name of the client
    Returns:
        str: File path of the generated Word document
    """
    doc = Document()
    # Title
    title = doc.add_heading(f"Proposal for {client_name}", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Date
    date_str = datetime.now().strftime("%B %d, %Y")
    date_paragraph = doc.add_paragraph(date_str)
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Add spacing
    doc.add_paragraph("")
    # Executive Summary
    doc.add_heading("Executive Summary", level=1)
    exec_par = doc.add_paragraph(sections.get('executive_summary', ''))
    exec_par.style = 'Normal'
    # Scope of Work
    doc.add_heading("Scope of Work", level=1)
    scope_par = doc.add_paragraph(sections.get('scope_of_work', ''))
    scope_par.style = 'Normal'
    # Pricing
    doc.add_heading("Pricing", level=1)
    price_par = doc.add_paragraph(sections.get('pricing', ''))
    price_par.style = 'Normal'
    # Set font for all paragraphs
    normal_style = doc.styles['Normal']
    font = normal_style.font# type: ignore
    font.name = 'Calibri'
    font.size = Pt(11)
    # Add some spacing
    for para in doc.paragraphs:
        para.paragraph_format.space_after = Pt(8)
        para.paragraph_format.line_spacing = 1.2
    # Save file
    safe_client = client_name.replace(' ', '_').replace('/', '_')
    filename = f"Proposal_{safe_client}_{datetime.now().strftime('%Y%m%d')}.docx"
    file_path = os.path.join(os.getcwd(), filename)
    doc.save(file_path)
    return file_path
