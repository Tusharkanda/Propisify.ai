from docx import Document
from docx.shared import Pt
from io import BytesIO
import re

# --- Paragraph formatter ---
def add_proposal_paragraph(doc, line):
    line = line.strip()

    # Skip Markdown horizontal rules like --- or ***
    if line in ['---', '***', '___']:
        return

    if line == "":
        # Add a small vertical gap (e.g. 3pt) for blank lines
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(3)
        return

    paragraph = doc.add_paragraph()

    # Process inline bold using **text**
    parts = re.split(r'(\*\*.*?\*\*)', line)
    for part in parts:
        run = paragraph.add_run(part.replace('**', '') if '**' in part else part)
        if part.startswith('**') and part.endswith('**'):
            run.bold = True

    # Detect likely headings
    is_heading = (
        re.match(r'^Phase\s+\d+[:.]', line) or
        re.match(r'^- Deliverables[:]*$', line.strip()) or
        line.istitle()
    )

    if is_heading:
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(12)
        paragraph.paragraph_format.space_after = Pt(4)  # More space after headings
    else:
        paragraph.paragraph_format.space_after = Pt(2)  # Less after regular lines


# --- Markdown table parser ---
def add_markdown_table_to_doc(doc, markdown_table_text):
    lines = markdown_table_text.strip().split("\n")
    
    # Filter valid table lines, remove separator lines like |---|
    lines = [line for line in lines if "|" in line and not re.match(r'^\|[-| ]+\|$', line)]

    if not lines:
        return

    # Convert to 2D array
    table_data = [ [cell.strip() for cell in line.strip('|').split('|')] for line in lines ]

    # Normalize row lengths by padding with empty strings
    max_cols = max(len(row) for row in table_data)
    for row in table_data:
        while len(row) < max_cols:
            row.append("")

    # Create table
    rows = len(table_data)
    table = doc.add_table(rows=rows, cols=max_cols)
    table.style = "Table Grid"

    for i, row in enumerate(table_data):
        for j, cell in enumerate(row):
            cell_paragraph = table.cell(i, j).paragraphs[0]
            cell_paragraph.clear()  # Clear default paragraph

            # Handle bold markdown in table like **text**
            parts = re.split(r'(\*\*.*?\*\*)', cell)
            for part in parts:
                run = cell_paragraph.add_run(part.replace('**', '') if '**' in part else part)
                if part.startswith('**') and part.endswith('**'):
                    run.bold = True

            # Bold the entire first row (header)
            if i == 0:
                for run in cell_paragraph.runs:
                    run.bold = True



# import re
# from docx.shared import Pt  # <--- Make sure this is imported!

def add_contract_paragraph(doc, line):
    # Detect headings: capitalized section titles or numbered sections
    is_heading = (
        line.strip().isalpha() and line.strip().istitle()
    ) or re.match(r'^\d+\.\s', line.strip())

    if line.strip() == "":
        doc.add_paragraph()
        return

    paragraph = doc.add_paragraph()
    run = paragraph.add_run(line.strip())

    if is_heading:
        run.bold = True
        run.font.size = Pt(12)  # Increase font size
        paragraph.paragraph_format.space_after = Pt(6)  # Add spacing below
