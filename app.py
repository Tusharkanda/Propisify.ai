import streamlit as st
from ocr_processor import OCRProcessor
from ai_generator import AIGenerator
from data_manager import DataManager
from document_parser import parse_proposal_sections, extract_metadata_from_text
from document_exporter import export_to_word
import os
from io import BytesIO
from docx import Document
import logging
import traceback
from datetime import datetime
from langchain_core.messages import HumanMessage, AIMessage
from langgraph.types import Command
from langchain_core.runnables.config import RunnableConfig
import uuid
from markdown_utils import add_proposal_paragraph, add_contract_paragraph, add_markdown_table_to_doc

from ai_agent import app

SUPPORTED_FILE_TYPES = [".pdf", ".png", ".jpg", ".jpeg", ".txt"]
INDUSTRY_OPTIONS = ["Technology", "Healthcare", "Finance", "Manufacturing", "Retail", "Other"]
SERVICE_OPTIONS = ["Consulting", "Software Development", "Data Analytics", "Marketing", "Other"]
TEAM_NAME = "BaagarBille"
TAGLINE = "AI-Powered Proposal Generation for Modern Business"

# Contract categories for dropdown
CONTRACT_CATEGORIES = [
    "AI Sales",
    "SaaS Healthcare",
    "Marketing",
    "Finance Analytics",
    "EdTech Platform",
    "Other"
]

# --- Logging Setup ---
logging.basicConfig(level=logging.INFO, format='%(asctime)s %(levelname)s %(message)s')
logger = logging.getLogger(__name__)

# --- ICONS ---
ICONS = {
    "upload": "📤",
    "generate": "🤖",
    "stats": "📊",
    "edit": "✏️",
    "export": "📄",
    "success": "✅",
    "error": "❌",
    "info": "ℹ️",
    "step": "🔹"
}


# STYLING THE TABLIST

st.markdown("""
    <style>
        div[data-baseweb="tab-list"] {
            background-color: #202020;
            border-radius: 12px;
            gap: 2px;
        }
    </style>
""", unsafe_allow_html=True)

# STYLING THE TAB BUTTONS

st.markdown("""
    <style>
        div[data-baseweb="tab-list"] button {
            height: 100%;
            padding: 1rem 1.5rem 1rem 1.5rem;
        }
        div[data-baseweb="tab-list"] button:hover {
            background-color: black;
            color: white;
        }
    </style>
""", unsafe_allow_html=True)

# STYLING ALL THE INPUT ELEMENTS

st.markdown("""
    <style>
        input[id^="text_input"], textarea[id^="text_area"] {
            background-color: #202020;
        }
            
        div[data-baseweb="select"] > div {
            background-color: #202020;
        }
        
        div[data-baseweb="base_input] > input{
            background-color: #202020;
        }
            
        #root {
            background-color: white;  
        }
            
    </style>
""", unsafe_allow_html=True)

def show_header():
    st.markdown(f"""
    <style>
        .main-header {{
            background: none;
            padding: 2rem;
            border-radius: 12px;
            border: 2px solid white;
            color: white;
        }}
        .main-header h1 {{
            font-size: 2.5rem;
            margin-bottom: 0.3rem;
        }}
        .main-header h4 {{
            margin-top: 0;
            color: #f0f0f0;
        }}
        .tagline {{
            font-style: italic;
            color: #ffc107;
        }}
    </style>
    <div class='main-header'>
        <h1>AI Proposal Writer</h1>
        <h4>{TAGLINE}</h4>
        <p class='tagline'> Hackathon Team: {TEAM_NAME}</p>
    </div>
    """, unsafe_allow_html=True)


def show_image():
    st.markdown("""
        <div class='random' style='display: flex; flex-direction: row; align-items: center; justify-content: center;">
    """, unsafe_allow_html=True)
    st.image('assets/images/logo.svg', width=350, caption='AI-Powered Proposal Writer', use_container_width=False)
    st.markdown("</div>", unsafe_allow_html=True)


def show_workflow():
    st.markdown(f"""
    <div style='margin-bottom:1.5rem;'>
        <h4>{ICONS['step']} <b>Step 1:</b> Upload proposals </h4>
        <h4>{ICONS['step']} <b>Step 2:</b> Generate a new proposal with AI </h4>
        <h4>{ICONS['step']} <b>Step 3:</b> Edit/export proposal</h4>
    </div>
    """, unsafe_allow_html=True)

def show_cards():
    st.markdown("""
        <style>
            .workflow-container {
                display: flex;
                justify-content: space-between;
                gap: 1.5rem;
                flex-wrap: wrap;
            }

            .workflow-card {
                flex: 1 1 30%;
                background-color: #1e1e1e;
                color: #f1f1f1;
                padding: 2rem;
                border-radius: 16px;
                box-shadow: 0 4px 20px rgba(0, 0, 0, 0.3);
                transition: transform 0.2s ease, box-shadow 0.2s ease;
                min-width: 250px;
            }

            .workflow-card:hover {
                transform: translateY(-5px);
                box-shadow: 0 6px 25px rgba(0, 0, 0, 0.5);
            }

            .circle {
                width: 40px;
                height: 40px;
                border-radius: 50%;
                display: flex;
                align-items: center;
                justify-content: center;
                font-weight: bold;
                margin-bottom: 1rem;
                font-size: 1.1rem;
            }

            .circle-1 {
                background-color: rgba(0, 255, 100, 0.1);
                color: #00ff64;
            }

            .circle-2 {
                background-color: rgba(255, 165, 0, 0.1);
                color: #ffa500;
            }

            .circle-3 {
                background-color: rgba(138, 43, 226, 0.1);
                color: #c084fc;
            }

            .card-title {
                font-size: 1.2rem;
                font-weight: bold;
                margin-bottom: 0.5rem;
            }

            .card-desc {
                font-size: 0.95rem;
                line-height: 1.4;
                color: #cfcfcf;
            }
        </style>

        <div class="workflow-container">
        <div class="workflow-card">
        <div class="circle circle-1">1</div>
        <div class="card-title">Data Collection</div>
        <div class="card-desc">
            Upload diverse datasets containing various types of information for analysis.
        </div>
        </div>
        <div class="workflow-card">
            <div class="circle circle-2">2</div>
            <div class="card-title">Model Training</div>
            <div class="card-desc">
                Allow the AI model to learn patterns and insights from the provided data. This may take some time.
            </div>
        </div>
        <div class="workflow-card">
        <div class="circle circle-3">3</div>
        <div class="card-title">Insights Delivery</div>
        <div class="card-desc">
            Receive comprehensive insights and analysis results based on the trained AI model.
        </div>
        </div>
        </div>
    """, unsafe_allow_html=True)
 
def detect_doc_type(text: str) -> str:
    """
    Detect if the text is a contract or proposal based on keywords/clauses.
    Returns 'contract' or 'proposal'.
    """
    contract_keywords = [
        'parties', 'term', 'scope of work', 'payment terms', 'confidentiality',
        'intellectual property', 'warranties', 'limitation of liability', 'termination',
        'dispute resolution', 'governing law', 'miscellaneous', 'agreement', 'contract'
    ]
    proposal_keywords = [
        'executive summary', 'scope of work', 'pricing', 'timeline', 'proposal'
    ]
    text_lower = text.lower()
    contract_hits = sum(1 for k in contract_keywords if k in text_lower)
    proposal_hits = sum(1 for k in proposal_keywords if k in text_lower)
    if contract_hits >= 3 and contract_hits > proposal_hits:
        return 'contract'
    return 'proposal'

def detect_contract_category(text: str) -> str:
    """
    Detect contract category based on keywords in the text.
    """
    mapping = {
        "AI Sales": ["ai", "artificial intelligence", "machine learning"],
        "SaaS Healthcare": ["saas", "healthcare", "ehr", "patient", "medical"],
        "Marketing": ["marketing", "campaign", "content creation", "digital marketing"],
        "Finance Analytics": ["bank", "finance", "analytics", "dashboard", "financial"],
        "EdTech Platform": ["education", "learning platform", "edtech", "online learning"]
    }
    text_lower = text.lower()
    for cat, keywords in mapping.items():
        if any(k in text_lower for k in keywords):
            return cat
    return "Other"

def main():
    st.set_page_config(page_title="AI Proposal Writer - Hackathon Demo", layout="wide")
    # show_header()
    # show_image()
    # show_workflow()
    st.markdown("""
                <style>
                    .logo-text {
                        font-family: 'Poppins', sans-serif;
                        font-size: 5rem;
                        font-weight: 800;
                
                        background: linear-gradient(90deg, #3B82F6, #22D3EE);
                        background-clip: text;
                        -webkit-background-clip: text;
                        color: transparent;
                        -webkit-text-fill-color: transparent;
                                        
                        letter-spacing: 1px;
                        text-shadow: 0 0 20px #C084FC66;
                        font-style: italic;
                        text-align: center;
                        margin-top: -75px;
                        margin-bottom: 20px;
                    }
                </style>
                <div class='logo-text'>
                    Proposify
                </div>
                """, unsafe_allow_html=True)
    show_cards()

    # --- SIDEBAR ---
    # Removed sidebar UI since API key is hardcoded
    # api_key = st.sidebar.text_input("Gemini API Key", type="password", help="Get your Gemini API key from Google AI Studio or Google Cloud Console.")
    # api_status = f"{ICONS['success']} Set" if api_key else f"{ICONS['error']} Not Set"
    # st.sidebar.markdown(f"**Gemini API Key Status:** {api_status}")
    # st.sidebar.info("All data is processed locally. No files are uploaded to any server.")
    # demo_mode = st.sidebar.checkbox("Demo Mode (use sample data)", value=False, help="Try the app with sample proposals if you don't have your own.")
    demo_mode = False
    api_key = True  # Placeholder, not used since key is hardcoded in ai_generator.py

    # --- SESSION STATE ---
    if 'data_manager' not in st.session_state:
        st.session_state['data_manager'] = DataManager()
        try:
            st.session_state['data_manager'].initialize_database()
        except Exception as e:
            logger.error(f"Database init error: {e}")
            st.error(f"Failed to initialize database: {e}")
    if 'processed_proposals' not in st.session_state:
        st.session_state['processed_proposals'] = []
    if 'generated_proposal' not in st.session_state:
        st.session_state['generated_proposal'] = None
    if 'generated_contract' not in st.session_state:
        st.session_state['generated_contract'] = None
    if 'thread_id' not in st.session_state:
        st.session_state['thread_id'] = str(uuid.uuid4())
    

    ocr_processor = OCRProcessor()

    # --- TABS ---
    
    st.markdown("""
    <style>
    .stTabs [data-baseweb="tab"] {{
        background-color: #f8f9fa;
        border: 1px solid #dee2e6;
        padding: 0.8rem;
        border-radius: 0.5rem 0.5rem 0 0;
        margin-right: 5px;
    }}
    .stTabs [aria-selected="true"] {{
        background-color: #ffffff;
        border-bottom: 2px solid #f1c40f;
        font-weight: 600;
    }}
    </style>
    """, unsafe_allow_html=True)

    tab_labels = [
        f"🗃️ Upload",
        f"⚙️ Generate",
        f"🖥️ Assistant"
    ]
    tabs = st.tabs(tab_labels)

    # --- UPLOAD TAB ---
    with tabs[0]:
        st.subheader(f"Upload SOWs or Contracts")
        st.markdown(f"<hr style='margin-top: -10px; border: 1px solid #ff4b4b;'>", unsafe_allow_html=True)
        uploaded_files = st.file_uploader(
            "Upload files (.pdf, .png, .jpg, .jpeg, .txt)",
            type=[ft.replace('.', '') for ft in SUPPORTED_FILE_TYPES],
            accept_multiple_files=True,
            help="You can upload multiple files at once."
        )
        if uploaded_files:
            progress_text = "Processing uploaded files..."
            progress_bar = st.progress(0, text=progress_text)
            total_files = len(uploaded_files)
            with st.spinner("Auto-processing files..."):
                for idx, file in enumerate(uploaded_files):
                    file_ext = os.path.splitext(file.name)[1].lower()
                    try:
                        # --- Extract text ---
                        if file_ext in [".png", ".jpg", ".jpeg"]:
                            extracted_text = ocr_processor.extract_text_from_image(file)
                        elif file_ext == ".pdf":
                            extracted_text = ocr_processor.extract_text_from_pdf(file)
                        elif file_ext == ".txt":
                            extracted_text = file.read().decode("utf-8", errors="ignore")
                        else:
                            st.warning(f"Unsupported file type: {file.name}")
                            continue
                        # --- Detect doc type ---
                        doc_type = detect_doc_type(extracted_text)
                        contract_category = detect_contract_category(extracted_text) if doc_type == 'contract' else None
                        # --- Parse sections & metadata ---
                        if doc_type == 'proposal':
                            sections = parse_proposal_sections(extracted_text)
                            auto_metadata = extract_metadata_from_text(extracted_text)
                        else:
                            sections = {}
                            auto_metadata = {"party_a": "", "party_b": ""}
                        meta = auto_metadata.copy()
                        meta['doc_type'] = doc_type
                        if contract_category:
                            meta['contract_category'] = contract_category
                        # --- Store in vector DB ---
                        success = st.session_state['data_manager'].store_proposal(
                            text=extracted_text,
                            metadata=meta
                        )
                        if success:
                            st.session_state['processed_proposals'].append({
                                "file": file.name,
                                "text": extracted_text,
                                "metadata": meta,
                                "sections": sections
                            })
                        else:
                            st.error(f"{ICONS['error']} Failed to store: {file.name}")
                    except Exception as e:
                        logger.error(f"Error processing {file.name}: {traceback.format_exc()}")
                        st.error(f"{ICONS['error']} Error processing {file.name}: {e}")
                    progress_bar.progress((idx + 1) / total_files, text=f"Processed {idx + 1} of {total_files} files...")
            progress_bar.empty()
            st.success(f"{ICONS['success']} All files processed and stored.")
        # Demo mode: show sample data if no docs
        if (demo_mode or not st.session_state['processed_proposals']):
            sample_text = """Executive Summary\nThis is a sample executive summary.\n\nScope of Work\nSample scope of work content.\n\nPricing\nSample pricing section."""
            sample_meta = {"client_name": "Demo Client", "industry": "Technology", "service_type": "Consulting", "doc_type": "proposal"}
            st.session_state['processed_proposals'] = [{
                "file": "sample.txt",
                "text": sample_text,
                "metadata": sample_meta,
                "sections": parse_proposal_sections(sample_text)
            }]

    # --- GENERATE TAB ---
    with tabs[1]:
        st.subheader(f"Generate New Document")
        st.markdown(f"<hr style='margin-top: -10px; border: 1px solid #ff4b4b;'>", unsafe_allow_html=True)
        doc_type_choice = st.radio("What do you want to generate?", ["Proposal", "Contract"], horizontal=True)
        
        if doc_type_choice == "Proposal":
            client_name = st.text_input("Client Name", key="gen_client_name", help="Enter the client's name.")
            industry = st.selectbox("Industry", INDUSTRY_OPTIONS, key="gen_industry", help="Select the industry.")
            service_type = st.selectbox("Service Type", SERVICE_OPTIONS, key="gen_service_type", help="Select the service type.")
            additional_reqs = st.text_area("Additional Requirements (optional)", key="gen_additional", help="Add any special requirements or notes.")
            generate_btn = st.button(f"{ICONS['generate']} Generate Proposal", disabled=not api_key)
            if generate_btn:
                if not api_key:
                    st.error(f"{ICONS['error']} Please enter your Gemini API key in the sidebar.")
                elif not client_name or not industry or not service_type:
                    st.error(f"{ICONS['error']} Please fill in all required fields.")
                else:
                    with st.spinner("Searching for similar proposals and generating content..."):
                        try:
                            similar = st.session_state['data_manager'].search_similar_proposals(
                                query=f"{client_name} {industry} {service_type}",
                                industry=industry if industry != "Other" else None,
                                service_type=service_type if service_type != "Other" else None
                            )
                            ai_gen = AIGenerator()
                            proposal = ai_gen.generate_full_proposal(
                                client_name=client_name,
                                industry=industry,
                                service_type=service_type,
                                similar_proposals=similar
                            )
                            if additional_reqs:
                                proposal['scope_of_work'] += f"\n\nAdditional Requirements:\n{additional_reqs}"
                            st.session_state['generated_proposal'] = proposal
                            st.success(f"{ICONS['success']} Proposal generated!")
                            # --- Display proposal ---
                            with st.expander(f"{ICONS['edit']} Executive Summary", expanded=True):
                                st.write(proposal['executive_summary'])
                            with st.expander(f"{ICONS['edit']} Scope of Work", expanded=True):
                                st.write(proposal['scope_of_work'])
                            with st.expander(f"{ICONS['edit']} Pricing Section", expanded=True):
                                st.write(proposal['pricing'])
                            # --- Edit Proposal ---
                            # st.markdown("---")
                            # st.subheader(f"{ICONS['edit']} Edit Proposal Sections")
                            # exec_edit = st.text_area("Edit Executive Summary", value=proposal['executive_summary'], key="edit_exec")
                            # scope_edit = st.text_area("Edit Scope of Work", value=proposal['scope_of_work'], key="edit_scope")
                            # price_edit = st.text_area("Edit Pricing Section", value=proposal['pricing'], key="edit_price")
                            # --- Export to Word ---
                            # if st.button(f"{ICONS['export']} Export to Word"):
                            #     try:
                            #         sections = {
                            #             'executive_summary': exec_edit,
                            #             'scope_of_work': scope_edit,
                            #             'pricing': price_edit
                            #         }
                            #         file_path = export_to_word(sections, client_name)
                            #         with open(file_path, "rb") as f:
                            #             st.download_button(
                            #                 label="Download Proposal as .docx",
                            #                 data=f,
                            #                 file_name=os.path.basename(file_path),
                            #                 mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            #             )
                            #     except Exception as e:
                            #         logger.error(f"Export error: {traceback.format_exc()}")
                            #         st.error(f"{ICONS['error']} Failed to export Word document: {e}")
                            
                            
                            # --- Download Proposal as .docx (LLM content only) ---
                            from docx import Document
                            from io import BytesIO


                            if st.session_state['generated_proposal']:
                                buffer = BytesIO()
                                doc = Document()
                                doc.add_heading('Proposal', 0)

                                # doc.add_paragraph(proposal['executive_summary'])
                                # doc.add_paragraph(proposal['scope_of_work'])
                                # doc.add_paragraph(proposal['pricing'])

                                # --- Executive Summary ---
                                doc.add_heading("Executive Summary", level=1)
                                for line in proposal["executive_summary"].splitlines():
                                    add_proposal_paragraph(doc, line)

                                # --- Scope of Work ---
                                doc.add_heading("Scope of Work", level=1)
                                for line in proposal["scope_of_work"].splitlines():
                                    add_proposal_paragraph(doc, line)

                                # --- Pricing Table ---
                                doc.add_heading("Pricing", level=1)
                                try:
                                    add_markdown_table_to_doc(doc, proposal["pricing"])
                                except Exception as e:
                                    doc.add_paragraph("Failed to parse pricing table. Raw:\n" + proposal["pricing"])


                                doc.save(buffer)
                                buffer.seek(0)

                                st.download_button(
                                    label="Download as .docx",
                                    data=buffer.getvalue(),
                                    file_name="proposal.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                                )
                                
                        except Exception as e:
                            logger.error(f"Proposal generation error: {traceback.format_exc()}")
                            if "rate limit" in str(e).lower():
                                st.error(f"{ICONS['error']} Gemini API rate limit reached. Please try again later.")
                            else:
                                st.error(f"{ICONS['error']} Error generating proposal: {e}")
        else:
            contract_category = st.selectbox("Contract Category", CONTRACT_CATEGORIES, key="gen_contract_category")
            party_a = st.text_input("Party A Name", key="gen_party_a")
            party_b = st.text_input("Party B Name", key="gen_party_b")
            effective_date = st.date_input("Effective Date", key="gen_effective_date")
            additional_reqs = st.text_area("Additional Requirements (optional)", key="gen_contract_additional")
            generate_btn = st.button(f"{ICONS['generate']} Generate Contract", disabled=not api_key)
            if generate_btn:
                # Search for similar contracts in selected category
                all_contracts = [p for p in st.session_state['processed_proposals'] if p['metadata'].get('doc_type') == 'contract' and p['metadata'].get('contract_category') == contract_category]
                ai_gen = AIGenerator()
                if all_contracts:
                    similar_contracts = all_contracts[:3]
                    contract = ai_gen.generate_full_contract(
                        party_a=party_a,
                        party_b=party_b,
                        effective_date=str(effective_date),
                        similar_contracts=similar_contracts,
                        additional_requirements=additional_reqs
                    )
                else:
                    contract = ai_gen.generate_contract_from_template(
                        party_a=party_a,
                        party_b=party_b,
                        effective_date=str(effective_date)
                    )
                st.session_state['generated_contract'] = contract
                # print(f"Generated contract: {contract[:100]}...")  # Log first 100 chars for debugging
                st.success(f"{ICONS['success']} Contract generated!")
                st.markdown(f"### Generated Contract\n\n" + contract)
                # --- Download Contract as .docx ---
                from docx import Document
                from io import BytesIO
                if st.session_state['generated_contract']:
                    contract_buffer = BytesIO()
                    doc = Document()
                    doc.add_heading('Contract', 0)

                    # doc.add_paragraph(contract)
                    for line in contract.splitlines():
                        add_contract_paragraph(doc, line)

                    doc.save(contract_buffer)
                    contract_buffer.seek(0)
                    st.download_button(
                        label="Download Contract as .docx",
                        data=contract_buffer.getvalue(),
                        file_name="contract.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

    # --- CHATS TAB ---
    with tabs[2]:
        # change here: 
        if st.session_state['generated_contract'] is None and st.session_state['generated_proposal'] is None:
            st.warning("Please generate a proposal or contract first in the 'Generate' tab.", icon="⚠️")
        else:
            doc_text = ""

            # print("GENERATED CONTRACT: \n\n", st.session_state['generated_contract'])
            # print("GENERATED PROPOSAL: \n\n", st.session_state['generated_proposal'])

            if (st.session_state['generated_contract'] is not None):
                doc_text = st.session_state['generated_contract']
            elif (st.session_state['generated_proposal'] is not None):
                for key, value in st.session_state['generated_proposal'].items():
                    doc_text += value

            human_message_content = (
                "I will give you a proposal or contract, and I want you to make changes according to my subsequent messages.\n"
                "The draft contract or "
                "proposal is:\n\n"
                + doc_text
            )
            
            st.subheader("💬 Chat History")
            if 'chat_history' not in st.session_state:
                st.session_state['chat_history'] = []
            
            config: RunnableConfig = RunnableConfig(configurable={"thread_id": st.session_state['thread_id']})
            
            if (len(st.session_state['chat_history']) == 0):
                first_human_message = HumanMessage(content=human_message_content)
                interrupt = app.invoke({'messages': [first_human_message]}, config=config)

                AIcontent = (
                    "Hello! I am your AI-powered business proposal and contract assistant. "
                    "I'm here to help you review, revise, and improve your proposals or contracts based on your feedback. "
                    "Please let me know what changes or enhancements you'd like to make, and I'll assist you every step of the way!"
                    "\n\nHere is the draft proposal or contract:\n\n"
                ) + doc_text

                st.session_state['chat_history'].append(AIMessage(
                    content=AIcontent
                ))
            
            # Display chat history
            for i, entry in enumerate(st.session_state['chat_history']):
                if isinstance(entry, HumanMessage) and i > 0:
                    st.markdown(f'<div style="text-align: right; font-size: 22px; font-weight: bolder;">User: </div> \n <div style="display: flex; justify-content: flex-end;"><div style="background-color: #4caf50; color: #fff; padding: 10px 15px; border-radius: 15px 15px 0 15px; max-width: 70%;"> {entry.content} </div></div>', unsafe_allow_html=True)
                elif isinstance(entry, AIMessage):
                    st.markdown(f'<div style="font-size: 22px; font-weight: bolder;">Bot: </div> \n <div style="display: flex; margin-bottom: 15px;"><div style="background-color: #3a3f58; color: #ffffff; padding: 10px 15px; border-radius: 15px 15px 15px 0; max-width: 70%;">{entry.content}</div></div>', unsafe_allow_html=True)
                st.markdown("---")

            # last_message = st.session_state['chat_history'][-1]
            # if isinstance(last_message, AIMessage):
            #     st.markdown(f"**Bot:** {last_message.content}")
            # else:
            #     st.markdown(f"**User:** {last_message.content}")
            # st.markdown("---")
            
            # Chat input
            user_input = st.text_input("Type your message...", key="chat_input")
            if st.button("Send", key="send_chat"):
                if user_input.strip():
                    result = app.invoke(
                        {'messages': st.session_state['chat_history'] + [HumanMessage(content=user_input)]},
                        config=config
                    )
                    st.session_state['chat_history'] = result['messages']
                    # print(f"\n\n\nChat history updated:\n {st.session_state['chat_history']}")
                    st.rerun()

if __name__ == "__main__":
    main()
